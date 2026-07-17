from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from pathlib import Path


class DocumentError(ValueError):
    pass


def read_text_guess(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "utf-16"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def markdown_to_text(markdown: str) -> str:
    text = markdown.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"```.*?```", "", text, flags=re.S)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"[*_~]{1,3}([^*_~]+)[*_~]{1,3}", r"\1", text)

    lines: list[str] = []
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"^>\s*", "", line)
        line = re.sub(r"^[-+*]\s+", "- ", line)
        line = re.sub(r"^\d+[.)]\s+", "", line)
        lines.append(line)
    return "\n".join(lines).strip()


def docx_to_text(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise DocumentError("缺少 python-docx 依赖，请先安装 requirements.txt。") from exc

    document = docx.Document(str(path))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())
        else:
            chunks.append("")
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append("    ".join(cells))
    return "\n".join(chunks).strip()


def convert_doc_to_docx(path: Path) -> Path:
    converter = shutil.which("soffice") or shutil.which("libreoffice")
    if not converter:
        raise DocumentError("当前机器没有找到 LibreOffice/soffice，暂不能直接读取 .doc。请先另存为 .docx。")

    out_dir = Path(tempfile.mkdtemp(prefix="handdraft-doc-"))
    command = [
        converter,
        "--headless",
        "--convert-to",
        "docx",
        "--outdir",
        str(out_dir),
        str(path),
    ]
    result = subprocess.run(command, capture_output=True, text=True, timeout=90, check=False)
    if result.returncode != 0:
        raise DocumentError(f".doc 转换失败：{result.stderr or result.stdout}".strip())
    converted = out_dir / f"{path.stem}.docx"
    if not converted.exists():
        matches = list(out_dir.glob("*.docx"))
        if not matches:
            raise DocumentError(".doc 转换后没有生成 .docx 文件。")
        converted = matches[0]
    return converted


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown"}:
        return markdown_to_text(read_text_guess(path))
    if suffix == ".txt":
        return read_text_guess(path).strip()
    if suffix == ".docx":
        return docx_to_text(path)
    if suffix == ".doc":
        return docx_to_text(convert_doc_to_docx(path))
    raise DocumentError("仅支持 .md、.markdown、.txt、.docx；.doc 需要本机安装 LibreOffice。")
