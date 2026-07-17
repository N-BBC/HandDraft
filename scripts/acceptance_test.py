from __future__ import annotations

import io
import json
import tempfile
import urllib.error
import urllib.request
import uuid
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw


ROOT = Path(__file__).resolve().parents[1]
BASE_URL = "http://127.0.0.1:8017"


def request_json(path: str, data: bytes | None = None, headers: dict[str, str] | None = None) -> dict:
    request = urllib.request.Request(BASE_URL + path, data=data, headers=headers or {})
    with urllib.request.urlopen(request, timeout=60) as response:
        assert response.status == 200, (path, response.status)
        return json.loads(response.read().decode("utf-8"))


def multipart(fields: dict[str, str], files: dict[str, Path]) -> tuple[bytes, str]:
    boundary = f"----HandDraft{uuid.uuid4().hex}"
    body = io.BytesIO()
    for name, value in fields.items():
        body.write(f"--{boundary}\r\n".encode())
        body.write(f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode())
        body.write(value.encode("utf-8"))
        body.write(b"\r\n")
    for name, path in files.items():
        body.write(f"--{boundary}\r\n".encode())
        body.write(
            f'Content-Disposition: form-data; name="{name}"; filename="{path.name}"\r\n'.encode()
        )
        body.write(b"Content-Type: application/octet-stream\r\n\r\n")
        body.write(path.read_bytes())
        body.write(b"\r\n")
    body.write(f"--{boundary}--\r\n".encode())
    return body.getvalue(), f"multipart/form-data; boundary={boundary}"


def render(document: Path, font_id: str, settings: dict, **assets: Path) -> dict:
    body, content_type = multipart(
        {"font_id": font_id, "settings": json.dumps(settings)},
        {"document": document, **assets},
    )
    result = request_json("/api/render", body, {"Content-Type": content_type})
    assert result["page_count"] >= 1
    output_paths = [page["url"] for page in result["pages"]]
    output_paths.extend([result["pdf_url"], result["zip_url"]])
    for output_path in output_paths:
        with urllib.request.urlopen(BASE_URL + output_path, timeout=30) as response:
            assert response.status == 200, output_path
            assert int(response.headers.get("Content-Length", "0")) > 500, output_path
    return result


def main() -> None:
    health = request_json("/api/health")
    assert health["ok"] is True
    fonts = request_json("/api/fonts")
    default_id = fonts["default_id"]
    default_font = next(font for font in fonts["fonts"] if font["id"] == default_id)
    reference_default = ROOT / "data" / "fonts" / "李国夫手写体.ttf"
    expected_default = reference_default.name if reference_default.exists() else "Xiaolai-Regular.ttf"
    assert default_font["filename"] == expected_default, default_font

    base_settings = {
        "scene_mode": "scan",
        "page_width": 640,
        "page_height": 900,
        "font_size": 26,
        "line_height": 42,
        "position_jitter": 1.2,
        "rotation_jitter": 1.4,
        "size_jitter": 1.1,
        "max_pages": 2,
    }
    sample = ROOT / "samples" / "sample.md"
    completed: list[str] = []
    for paper_style in (
        "photo_blank",
        "reference_blank_desk",
        "reference_blank_warm",
        "reference_blank_cool",
        "reference_blank_mono",
        "reference_lined_photo",
        "reference_lined_clean",
        "reference_notebook",
        "reference_report_body",
        "reference_report_cover",
        "grid",
        "blank",
    ):
        render(sample, default_id, {**base_settings, "paper_style": paper_style})
        completed.append(f"paper:{paper_style}")

    with tempfile.TemporaryDirectory(prefix="handdraft-acceptance-") as temp_dir:
        temp = Path(temp_dir)
        docx_path = temp / "sample.docx"
        document = Document()
        document.add_heading("HandDraft Word 测试", level=1)
        document.add_paragraph("这是一段正常手写字体的 DOCX 转换测试。")
        document.save(docx_path)
        render(docx_path, default_id, {**base_settings, "paper_style": "lined"})
        completed.append("document:docx")

        paper_path = temp / "custom-paper.png"
        paper = Image.new("RGB", (720, 1018), (248, 245, 235))
        draw = ImageDraw.Draw(paper)
        for y in range(54, paper.height, 48):
            draw.line((0, y, paper.width, y), fill=(210, 221, 219), width=1)
        paper.save(paper_path)

        custom_result = render(
            sample,
            default_id,
            {
                **base_settings,
                "scene_mode": "desk_photo",
                "use_background_size": True,
                "direct_background": True,
                "paper_style": "custom_photo",
            },
            background=paper_path,
        )
        custom_page_url = custom_result["pages"][0]["url"]
        with urllib.request.urlopen(BASE_URL + custom_page_url, timeout=30) as response:
            custom_page = Image.open(io.BytesIO(response.read()))
            assert custom_page.size == paper.size, custom_page.size
        completed.extend(["background:single-photo", "scene:direct-template"])

    secret = "sk-acceptance-secret-1234567890"
    payload = json.dumps({"provider": "custom", "model": "test", "apiKey": secret}).encode("utf-8")
    ai_result = request_json("/api/ai/glyph-kit", payload, {"Content-Type": "application/json"})
    assert secret not in json.dumps(ai_result)
    assert ai_result["key_hint"] != secret
    completed.append("security:api-key-redacted")

    print(f"acceptance test passed ({len(completed)} checks)")
    print(", ".join(completed))


if __name__ == "__main__":
    try:
        main()
    except urllib.error.URLError as exc:
        raise SystemExit(f"HandDraft server is not reachable at {BASE_URL}: {exc}") from exc
