from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from PIL import Image

from handdraft.document import docx_to_text, markdown_to_text
from handdraft.fonts import find_default_font
from handdraft.open_fonts import OPEN_FONTS
from handdraft.renderer import RenderSettings, render_pages, save_outputs
from handdraft.security import redact_secret


class HandDraftCoreTests(unittest.TestCase):
    def test_markdown_conversion(self) -> None:
        result = markdown_to_text("# Title\n\n- first\n- [second](https://example.com)")
        self.assertIn("Title", result)
        self.assertIn("- first", result)
        self.assertIn("second", result)
        self.assertNotIn("https://", result)

    def test_docx_conversion(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "input.docx"
            document = Document()
            document.add_heading("Word test", level=1)
            document.add_paragraph("Chinese content 123")
            document.save(path)
            result = docx_to_text(path)
        self.assertIn("Word test", result)
        self.assertIn("Chinese content 123", result)

    def test_render_png_pdf_and_zip(self) -> None:
        font = find_default_font()
        self.assertIsNotNone(font)
        settings = RenderSettings(
            page_width=640,
            page_height=900,
            margin_left=50,
            margin_right=50,
            margin_top=60,
            margin_bottom=60,
            font_size=26,
            line_height=42,
            scene_mode="scan",
            max_pages=2,
            seed=9,
        )
        pages = render_pages("HandDraft render test\nSecond line", font, None, settings)
        with tempfile.TemporaryDirectory() as temp_dir:
            outputs = save_outputs(pages, Path(temp_dir))
            self.assertTrue((Path(temp_dir) / outputs["page_files"][0]).exists())
            self.assertGreater((Path(temp_dir) / outputs["pdf"]).stat().st_size, 1000)
            self.assertGreater((Path(temp_dir) / outputs["zip"]).stat().st_size, 1000)

    def test_direct_background_preserves_size_and_margins(self) -> None:
        font = find_default_font()
        self.assertIsNotNone(font)
        with tempfile.TemporaryDirectory() as temp_dir:
            background = Path(temp_dir) / "custom-background.png"
            Image.new("RGB", (700, 1000), (245, 242, 232)).save(background)
            settings = RenderSettings(
                margin_left=37,
                margin_right=83,
                margin_top=64,
                margin_bottom=70,
                font_size=26,
                line_height=42,
                scene_mode="desk_photo",
                direct_background=True,
                paper_style="custom_photo",
                max_pages=2,
            )
            pages = render_pages("自定义背景测试", font, background, settings)

        self.assertEqual(pages[0].size, (700, 1000))
        self.assertEqual(settings.margin_left, 37)
        self.assertEqual(settings.margin_right, 83)

    def test_template_keeps_user_spacing(self) -> None:
        font = find_default_font()
        self.assertIsNotNone(font)
        settings = RenderSettings(
            page_width=640,
            margin_left=31,
            margin_right=79,
            margin_top=60,
            margin_bottom=60,
            font_size=24,
            line_height=40,
            scene_mode="scan",
            paper_style="reference_blank_warm",
            max_pages=2,
        )
        pages = render_pages("左右间距测试", font, None, settings)

        self.assertEqual(pages[0].width, 640)
        self.assertEqual(settings.margin_left, 31)
        self.assertEqual(settings.margin_right, 79)

    def test_api_key_redaction(self) -> None:
        secret = "sk-testsecret1234567890"
        redacted = redact_secret(secret)
        self.assertNotIn("testsecret", redacted)
        self.assertTrue(redacted.startswith("sk-t"))
        self.assertTrue(redacted.endswith("7890"))

    def test_normal_open_font_catalog(self) -> None:
        normal = {font["family"] for font in OPEN_FONTS if font["category"] == "normal"}
        self.assertEqual(normal, {"Xiaolai", "LXGW WenKai"})


if __name__ == "__main__":
    unittest.main()
